from docx import Document

# File name
docx_path = "AI_ML_24_Month_Roadmap.docx"

# Create document
doc = Document()
doc.add_heading("AI & ML 24-Month Roadmap (6 hrs/week)", 0)

# Intro
doc.add_paragraph(
    "This roadmap is designed for a B.Tech student with 2 years left, "
    "who can dedicate ~6 hours per week and wants to be job-ready in AI/ML by graduation."
)

# Roadmap content
roadmap = [
    ("Months 1–4 (Foundations)", [
        "Python basics (syntax, loops, OOP, functions, libraries)",
        "Maths basics (linear algebra, probability, statistics, derivatives)",
        "SQL basics (SELECT, WHERE, JOIN, GROUP BY, aggregations)",
        "Libraries: NumPy, Pandas, Matplotlib",
        "Projects: Data cleaning + visualization, MySQL student DB"
    ]),
    ("Months 5–8 (Core Machine Learning)", [
        "Regression, classification, clustering, decision trees, random forest, SVM",
        "Model evaluation metrics (accuracy, precision, recall, F1, confusion matrix)",
        "Library: scikit-learn",
        "Projects: House price prediction, Spam detection, Customer segmentation"
    ]),
    ("Months 9–12 (Deep Learning Basics)", [
        "Neural networks (forward pass, backpropagation, activation, loss, optimizers)",
        "CNN basics (image classification)",
        "RNN & LSTM intuition",
        "Projects: MNIST digit classifier, Cat vs Dog image classifier"
    ]),
    ("Months 13–16 (Specialization + DSA start)", [
        "Choose specialization: Computer Vision (CV) or NLP",
        "Start DSA basics: arrays, strings, hashing, 2 pointers, binary search",
        "Projects: Transfer learning (CV) or Sentiment analysis (NLP)"
    ]),
    ("Months 17–20 (MLOps + Deployment)", [
        "Flask/FastAPI basics, Docker basics",
        "SQL integration with ML pipeline",
        "Cloud basics (Heroku free deploy, optional AWS/GCP)",
        "Projects: Deploy image classifier, End-to-end ML pipeline with MySQL"
    ]),
    ("Months 21–24 (Interview Prep + Portfolio polish)", [
        "DSA: DP, graphs, recursion, complexity analysis",
        "ML interview prep: bias-variance, overfitting, pipelines, feature engineering",
        "Polish Resume, LinkedIn, GitHub",
        "Final Projects: Object detection (CV) or Fake news detection (NLP)"
    ])
]

# Add roadmap content
for section, items in roadmap:
    doc.add_heading(section, level=1)
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

# Weekly routine
doc.add_heading("Weekly Study Routine (6 hrs/week)", level=1)
weekly = [
    "2 hrs → Theory & tutorials (ML/DL/SQL/Python)",
    "2 hrs → Project coding / Kaggle practice",
    "1 hr → DSA (LeetCode easy → medium)",
    "1 hr → Portfolio (GitHub updates, LinkedIn posts, resume polishing)"
]
for item in weekly:
    doc.add_paragraph(item, style="List Bullet")

# Save docx
doc.save(docx_path)
print(f"Saved: {docx_path}")